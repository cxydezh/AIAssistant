import os
from docx import Document
from docx.shared import Pt
import re
from collections import defaultdict
from docx.text.paragraph import Paragraph

def split_document_by_headings(input_path, output_folder):
    # 创建输出文件夹
    os.makedirs(output_folder, exist_ok=True)
    
    # 加载文档
    doc = Document(input_path)
    
    # 存储标题及其出现次数
    title_count = defaultdict(int)
    current_title = None
    current_content = []
    subdocuments = []
    
    # 支持的标题样式（中英文）
    heading_styles = [
        'Heading 1', 'Heading1', '标题 1', '一级标题', 'Title 1', 'heading 1', 'heading1', 'title 1',
        'HEADING 1', 'TITLE 1', '标题1', '一级', 'Heading', '标题', 'Title'
    ]
    
    print('--- 段落样式调试输出 ---')
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name.strip() if para.style and para.style.name else ''
        print(f"段落{idx}: '{text[:30]}' 样式: '{style_name}'")
    print('--- 结束 ---')
    
    # 遍历文档中的所有段落
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name.strip() if para.style and para.style.name else ''
        is_heading = style_name in heading_styles
        
        # 如果是标题
        if is_heading:
            print(f"检测到标题: '{text}' 样式: '{style_name}' (段落{idx})")
            # 保存前一个子文档（只要有current_title，不要求current_content非空）
            if current_title is not None:
                clean_title = re.sub(r'[\\/*?:"<>|]', "", current_title)
                title_count[clean_title] += 1
                count = title_count[clean_title]
                
                # 添加序号如果标题重复
                filename = f"{clean_title}_{count}.docx" if count > 1 else f"{clean_title}.docx"
                subdocuments.append((filename, current_content))
            
            # 开始新的子文档
            current_title = text
            current_content = [para]
        else:
            # 添加到当前内容
            if current_title is not None:
                current_content.append(para)
    
    # 保存最后一个子文档（只要有current_title）
    if current_title is not None:
        clean_title = re.sub(r'[\\/*?:"<>|]', "", current_title)
        title_count[clean_title] += 1
        count = title_count[clean_title]
        filename = f"{clean_title}_{count}.docx" if count > 1 else f"{clean_title}.docx"
        subdocuments.append((filename, current_content))
    
    print(f"共检测到{len(subdocuments)}个子文档")
    # 保存所有子文档
    for filename, content in subdocuments:
        output_path = os.path.join(output_folder, filename)
        new_doc = Document()
        
        # 复制内容到新文档
        for element in content:
            if isinstance(element, Paragraph):
                new_para = new_doc.add_paragraph()
                for run in element.runs:
                    new_run = new_para.add_run(run.text)
                    # 复制样式
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
        
        new_doc.save(output_path)
        print(f"已保存: {output_path}")

if __name__ == "__main__":
    input_docx = r"D:\CursorCode\pythoncode\AIAssistant\AI辅助示例数据\6442477_陈某某_3\备份\病历.docx"
    output_dir = r"D:\CursorCode\pythoncode\AIAssistant\AI辅助示例数据\6442477_陈某某_3\病历资料"
    
    split_document_by_headings(input_docx, output_dir)
    print("文档拆分完成！")